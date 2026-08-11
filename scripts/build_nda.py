#!/usr/bin/env python3
"""
build_nda — Phase 4.5.

Generates docs/NDA_es-MX.docx (Tester NDA) from the canonical text below.
Spanish-first, single-signature, references LFPDPPP. Output is the binary
.docx the operator emails / WhatsApps to each tester before they receive
their invite code.

Usage:
    python3 scripts/build_nda.py
    python3 scripts/build_nda.py --out custom.docx

The text is in this script (not a separate template) so:
  1. Every change goes through PR review.
  2. The hash that gets logged in the audit trail (P4.2 consent flow)
     stays anchored to a specific git commit.
  3. The structural sanity gate validates the headings stay in place.
"""
from __future__ import annotations

import argparse
import hashlib
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_OUT = ROOT / "docs" / "NDA_es-MX.docx"

NDA_VERSION = "nda-v1"

# ============================================================================
# Canonical NDA text
# ============================================================================
TITLE = "ACUERDO DE CONFIDENCIALIDAD PARA PROBADOR DE BETA"
SUBTITLE = "Casa·Orquesta · Programa de pruebas cerradas · Junio 2026"

PARTIES_HEADING = "Partes"
PARTIES = (
    "Por una parte, Casa·Orquesta S.A. de C.V., con domicilio en Ciudad de "
    "México, México, representada por su apoderado legal "
    "(\"Casa·Orquesta\" o \"la Empresa\"), y por la otra, la persona física "
    "que firma al final de este documento (el \"Probador\")."
)

RECITALS_HEADING = "Antecedentes"
RECITALS = (
    "I.   Casa·Orquesta está desarrollando una aplicación móvil de asesoría "
    "inmobiliaria por voz para CDMX y Morelos, actualmente en beta cerrada.\n"
    "II.  La Empresa desea invitar al Probador a usar la aplicación en su "
    "fase pre-lanzamiento y recibir retroalimentación honesta y constructiva.\n"
    "III. Durante esta participación el Probador tendrá acceso a información "
    "confidencial de la Empresa que no es del dominio público."
)

CLAUSES = [
    (
        "1. Objeto",
        "El presente acuerdo regula la confidencialidad de toda información "
        "que el Probador reciba o conozca de Casa·Orquesta con motivo de su "
        "participación en el programa de beta cerrada, incluyendo —sin "
        "limitarse a— funcionalidades no lanzadas, estrategia de producto, "
        "métricas internas, materiales gráficos, listados de propiedades "
        "ficticias o reales, y conversaciones con el equipo."
    ),
    (
        "2. Definición de información confidencial",
        "Se considerará Información Confidencial cualquier dato que: (a) "
        "esté marcado como tal; (b) razonablemente se entienda como "
        "confidencial por su naturaleza; o (c) haya sido revelado bajo "
        "expectativa de confidencialidad. No incluye información que sea o "
        "llegue a ser de dominio público sin culpa del Probador."
    ),
    (
        "3. Obligaciones del Probador",
        "El Probador se compromete a: (a) usar la Información Confidencial "
        "únicamente para el propósito de probar la aplicación; (b) no "
        "compartirla con terceros sin consentimiento escrito de "
        "Casa·Orquesta; (c) no publicar capturas de pantalla, videos, o "
        "descripciones de la aplicación en redes sociales o medios públicos "
        "hasta el lanzamiento oficial; (d) reportar de buena fe los errores "
        "que encuentre a través de los canales designados."
    ),
    (
        "4. Vigencia",
        "Este acuerdo entra en vigor en la fecha de firma y permanece "
        "vigente por un (1) año a partir del lanzamiento público de la "
        "aplicación, o por dos (2) años a partir de la firma, lo que "
        "ocurra primero."
    ),
    (
        "5. Protección de datos personales (LFPDPPP)",
        "Casa·Orquesta tratará los datos personales del Probador (nombre, "
        "teléfono, voz, transcripciones y uso de la aplicación) conforme a "
        "la Ley Federal de Protección de Datos Personales en Posesión de "
        "los Particulares (LFPDPPP) y a su Aviso de Privacidad integral "
        "publicado en https://casaorquesta.mx/aviso-de-privacidad. El "
        "Probador podrá ejercer sus derechos ARCO (Acceso, Rectificación, "
        "Cancelación, Oposición) escribiendo a privacidad@casaorquesta.mx."
    ),
    (
        "6. Devolución / destrucción de materiales",
        "Al término de este acuerdo, el Probador deberá desinstalar la "
        "aplicación y destruir cualquier material físico o digital que "
        "haya recibido de Casa·Orquesta. La Empresa eliminará los datos "
        "personales del Probador conforme al artículo 11 de la LFPDPPP "
        "salvo las obligaciones legales contables (CFDI 4.0)."
    ),
    (
        "7. Incumplimiento y reparación",
        "El Probador reconoce que el incumplimiento de este acuerdo puede "
        "causar daños difíciles de cuantificar. Las partes acuerdan que "
        "Casa·Orquesta podrá solicitar medidas precautorias y "
        "compensación por daños comprobados ante los tribunales "
        "competentes de la Ciudad de México."
    ),
    (
        "8. Independencia del acuerdo",
        "Este documento no crea relación laboral, asociativa, ni "
        "representativa entre las partes. El Probador participa por "
        "voluntad propia y sin remuneración."
    ),
    (
        "9. Modificaciones",
        "Cualquier modificación al presente acuerdo deberá ser por escrito "
        "y firmada por ambas partes."
    ),
    (
        "10. Jurisdicción y ley aplicable",
        "Este acuerdo se rige por las leyes vigentes en los Estados Unidos "
        "Mexicanos. Las partes se someten a la jurisdicción de los "
        "tribunales competentes de la Ciudad de México, renunciando a "
        "cualquier otro fuero que pudiera corresponderles por razón de "
        "domicilio presente o futuro."
    ),
]

SIGNATURE_NOTE = (
    "Leído y aceptado en su totalidad. Al firmar, el Probador manifiesta "
    "que ha tenido la oportunidad de leer cada cláusula, formular preguntas "
    "y consultar con un asesor legal si así lo deseó."
)


# ============================================================================
# Builder
# ============================================================================
def _set_doc_defaults(doc: Document) -> None:
    """Make the default style Helvetica 11pt. Apple/Mac default font and
    a reasonable size for legal docs."""
    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(11)
    # Also tell Word's east-asian / hi-ansi fallback so the font sticks.
    rPr = style.element.get_or_add_rPr()
    for tag in ("w:rFonts",):
        existing = rPr.find(qn(tag))
        if existing is not None:
            rPr.remove(existing)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Helvetica")
    rFonts.set(qn("w:hAnsi"), "Helvetica")
    rFonts.set(qn("w:cs"), "Helvetica")
    rPr.append(rFonts)


def _set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)


def _add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"Casa·Orquesta S.A. de C.V. · NDA Beta · {NDA_VERSION}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _add_title_block(doc: Document) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(16)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run(SUBTITLE)
    rs.italic = True
    rs.font.size = Pt(11)
    rs.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Decorative rule
    rule = doc.add_paragraph()
    pPr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D4A24C")  # brand gold
    pbdr.append(bottom)
    pPr.append(pbdr)

    doc.add_paragraph()  # spacing


def _add_heading(doc: Document, text: str, *, size: int = 13) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def _add_body(doc: Document, text: str) -> None:
    for line in text.split("\n"):
        if not line.strip():
            doc.add_paragraph()
            continue
        p = doc.add_paragraph(line.strip())
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.3


def _add_signature_block(doc: Document) -> None:
    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run(SIGNATURE_NOTE)
    nr.italic = True
    nr.font.size = Pt(10)
    nr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()

    # Two-column signature table without visible borders.
    table = doc.add_table(rows=4, cols=2)
    table.autofit = True
    # Hide borders
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)

    cells = [
        ("Casa·Orquesta S.A. de C.V.", "El Probador"),
        ("", ""),
        ("_____________________________________", "_____________________________________"),
        ("Nombre / Cargo", "Nombre"),
    ]
    for r_idx, (left, right) in enumerate(cells):
        c1, c2 = table.rows[r_idx].cells
        c1.text = left
        c2.text = right
        for cell in (c1, c2):
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date.add_run("Fecha: ________ / ________ / 2026")
    dr.font.size = Pt(10)


def build(out: Path) -> bytes:
    doc = Document()
    _set_doc_defaults(doc)
    _set_margins(doc)
    _add_footer(doc)

    _add_title_block(doc)

    _add_heading(doc, PARTIES_HEADING)
    _add_body(doc, PARTIES)

    _add_heading(doc, RECITALS_HEADING)
    _add_body(doc, RECITALS)

    _add_heading(doc, "Cláusulas", size=14)
    for heading, body in CLAUSES:
        _add_heading(doc, heading, size=12)
        _add_body(doc, body)

    _add_signature_block(doc)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out.read_bytes()


def main(argv: list[str] | None = None) -> int:
    p = argparse.ArgumentParser(description="Build docs/NDA_es-MX.docx")
    p.add_argument("--out", type=Path, default=DEFAULT_OUT)
    p.add_argument("--print-hash", action="store_true",
                   help="Print SHA-256 of the resulting file.")
    args = p.parse_args(argv)
    blob = build(args.out)
    print(f"  ✓ wrote {args.out} ({len(blob):,} bytes)")
    if args.print_hash:
        print(f"  sha256: {hashlib.sha256(blob).hexdigest()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
